import random
import sys
from tqdm import tqdm
import requests
import time
from requests.utils import dict_from_cookiejar, cookiejar_from_dict
import json
import math
import re
import string
from urllib import parse
from docx import Document
import os
from docx.oxml.ns import qn
from docx.shared import RGBColor


class mooc_login():

    def __init__(self):
        self.session = requests.session()
        self.session.headers = {
            'referer': 'https://www.icourse163.org/',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.190 Safari/537.36',

        }

    def get_codeUrl_pollKey(self):
        self.session.get("https://www.icourse163.org/")
        url = "https://www.icourse163.org/logonByQRCode/code.do?width=182&height=182"
        response = self.session.get(url=url)
        dic = response.json()
        if dic["code"] == 0:
            print("获取codeUrl与pollKey成功")
        else:
            print("获取codeUrl与pollKey失败")
        codeUrl = dic["result"]["codeUrl"]
        pollKey = dic["result"]["pollKey"]
        return codeUrl, pollKey

    def show_QR_img(self, codeUrl):
        response = self.session.get(url=codeUrl)
        content = response.content
        with open('core\QR.png', 'wb') as f:
            f.write(content)
        os.startfile('core\QR.png')

    def get_status(self, pollKey):
        url = 'https://www.icourse163.org/logonByQRCode/poll.do?pollKey={}'.format(pollKey)
        while True:
            response = self.session.get(url=url)
            dic = response.json()
            if dic["result"]["codeStatus"] == 0:
                print("请及时扫码")
            elif dic["result"]["codeStatus"] == 1:
                print("请点击确认")
            else:
                print("登录成功")
                token = dic["result"]["token"]
                return token
            time.sleep(1)

    def save_cookie(self):
        codeUrl, pollKey = self.get_codeUrl_pollKey()
        self.show_QR_img(codeUrl)
        token = self.get_status(pollKey)
        params = {
            "token": token,
            "returnUrl": "aHR0cHM6Ly93d3cuaWNvdXJzZTE2My5vcmcv",
        }
        url = "https://www.icourse163.org/passport/logingate/mocMobChangeCookie.htm"
        self.session.get(url=url, params=params)
        self.session.get("https://www.icourse163.org/")
        cookie_jar = self.session.cookies
        cookie_dic = dict_from_cookiejar(cookie_jar)
        with open('core/cookies.json', 'w') as f:
            f.write(json.dumps(cookie_dic))
        print("cookie已保存")

    def check_cookie(self):
        print('正在检查cookie有效性')
        url = 'https://www.icourse163.org/'
        response = self.session.get(url=url)
        text = response.text
        if "个人中心" in text:
            print("cookie有效")
            return True
        else:
            print("cookie无效")
            return False

    def reade_cookie(self):
        b = os.path.exists("core/cookies.json")
        if b == False:
            print("没有cookie文件")
            return False
        else:
            with open(r'core/cookies.json', 'r') as f:
                cookie_dic = json.loads(f.read())
            cookie_jar = cookiejar_from_dict(cookie_dic)
            self.session.cookies = cookie_jar
            print("cookie已加载")
            return True

    def login(self):
        b = os.path.exists('core')
        if b == False:
            os.mkdir('core')
        b = self.reade_cookie()
        if b == True:
            b1 = self.check_cookie()
            if b1 == False:
                # 删除存在的失效cookie
                self.session.cookies.clear()
                self.save_cookie()
        else:
            self.save_cookie()

        self.session.headers = {
            "sec-ch-ua": '"Google Chrome";v="89", "Chromium";v="89", ";Not A Brand";v="99"',
            "sec-ch-ua-mobile": "?0",
            "sec-fetch-dest": "document",
            "sec-fetch-mode": "navigate",
            "sec-fetch-site": "same-origin",
            "upgrade-insecure-requests": "1",
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.72 Safari/537.3",
        }
        return self.session


class mooc_spider():

    def __init__(self):
        self.session = mooc_login().login()

    def get_csrfKey_userId(self):
        self.userId = dict_from_cookiejar(self.session.cookies)["NETEASE_WDA_UID"].split('#')[0]
        params = {
            "userId": self.userId
        }
        self.session.get('https://www.icourse163.org/home.htm', params=params)
        self.csrfKey = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        self.session.headers.update({"edu-script-token": self.csrfKey})
        return

    def get_pagesize(self):
        url = "https://www.icourse163.org/web/j/learnerCourseRpcBean.getPersonalLearningStatisticDto.rpc"
        data = {
            "uid": self.userId
        }
        params = {
            "csrfKey": self.csrfKey
        }
        response = self.session.post(url=url, params=params, data=data)
        learingCoursesCount = response.json()["result"]["learingCoursesCount"]
        pagesize = math.ceil(learingCoursesCount / 8) + 1
        return pagesize

    def get_courses(self, pagessize):
        courses_dic = {}
        url = "https://www.icourse163.org/web/j/learnerCourseRpcBean.getMyLearnedCoursePanelList.rpc"
        params = {
            "csrfKey": self.csrfKey
        }
        n_key = 0
        for p in range(1, pagessize):
            data = {
                "type": "30",
                "p": p,
                "psize": pagessize,
                "courseType": "1",
            }
            response = self.session.post(url, data=data, params=params)
            result = response.json()["result"]["result"]
            for i in result:
                name = i["name"]
                courses_id = i["termPanel"]["id"]
                school_name = i["schoolPanel"]["name"]
                list = [name, school_name, courses_id]
                courses_dic[str(n_key)] = list
                n_key += 1
        return courses_dic

    def get_timestamp(self):
        timestamp = int(time.time() * 1000)
        return str(timestamp)

    # httpSessionId
    def get_NTESSTUDYSI(self):
        NTESSTUDYSI = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        return NTESSTUDYSI

    # 获取测验相关信息
    def get_test_info(self, courses_id):
        test_name_id_dic = {}
        url = "https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLastLearnedMocTermDto.dwr"
        data = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'CourseBean',
            'c0-methodName': 'getLastLearnedMocTermDto',
            'c0-id': '0',
            'c0-param0': 'number:{}'.format(courses_id),
            'batchId': self.get_timestamp()
        }
        response = self.session.post(url=url, data=data)
        text = response.text.encode('utf8').decode('unicode-escape')
        option_text_list = re.findall('type.*?allowUpload|type.*?dwr\.engine', text, re.S)
        id_name_list = re.findall(
            's\d{1,3}\.id=(\d+);s\d{1,3}\.name="(.*?测试.*?)";|s\d{1,3}\.id=(\d+);s\d{1,3}\.name="(.*?测验.*?)";', text)
        for id_name in id_name_list:
            if id_name[0] != '' and id_name[1] != '':
                test_id = id_name[0]
                test_name = id_name[1]
            else:
                test_id = id_name[2]
                test_name = id_name[3]
            test_name_id_dic[test_name] = test_id
        return test_name_id_dic

    def submit_paper(self, test_id):
        data = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'MocQuizBean',
            'c0-methodName': 'getQuizPaperDto',
            'c0-id': '0',
            'c0-param0': 'string:{}'.format(test_id),
            'c0-param1': 'number:0',
            'c0-param2': 'boolean:false',
            'batchId': self.get_timestamp(),
        }
        url = "https://www.icourse163.org/dwr/call/plaincall/MocQuizBean.getQuizPaperDto.dwr"
        response = self.session.post(url=url, data=data)
        text = response.text.encode('utf-8').decode('unicode-escape')
        time.sleep(1)

        option_text_list = re.findall('type.*?allowUpload|type.*?dwr\.engine', text, re.S)
        content_id_list = []
        for option_text in option_text_list:
            content_id = re.findall('content="(.*?)";s\d+\.id=(\d+);', option_text, re.S)
            content_id_list.append(content_id)
        option_dict = {'analyse': '', 'answer': '', 'content': '', 'id': '', 'selectCount': ''}
        aid, tid, tname, end_type = \
            re.findall('dwr.engine.*?aid:(\d+),.*?tid:(\d+),tname:"(.*?)",type:(\d+)}', text, re.S)[0]
        submit_data = {
            "callCount": '1',
            "scriptSessionId": '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'MocQuizBean',
            'c0-methodName': 'submitAnswers',
            'c0-id': '0',
            "c0-e1": "number:{}".format(aid),
            "c0-e2": "null:null",
            "c0-e3": "boolean:false",
            "c0-e4": "null:null",
        }
        list_0 = re.findall("s\d+\.allowUpload.*?type=\d+;", text, re.S)
        id_list = []
        plainTextTitle_list = []
        score_list = []
        title_list = []
        type_list = []
        position_list = []
        for i in list_0:
            id = re.findall("s\d+\.id=(.*?);", i)
            id_list.extend(id)

            plainTextTitle = re.findall('s\d+\.plainTextTitle="(.*?)";', i)
            plainTextTitle_list.extend(plainTextTitle)

            score = re.findall('s\d+\.score=(.*?);', i)
            score_list.extend(score)

            title = re.findall('s\d+\.title="(.*?)";s\d+\.titleAttachment', i)
            title_list.extend(title)

            type = re.findall('s\d+\.type=(.*?);', i)
            type_list.extend(type)

            position = re.findall('s\d+\.position=(.*?);', i)
            position_list.extend(position)

        Object_Object_dict = {"allowUpload": "", 'analyse': '', 'description': '', 'fillblankType': '', 'gmtCreate': '',
                              'gmtModified': '', 'id': '', 'judgeDtos': '', 'judgerules': '', 'ojCases': '',
                              'ojMemLimit': '',
                              'ojNeedInput': '', 'ojSupportedLanguage': '', 'ojSupportedLanguageList': '',
                              'ojTimeLimit': '',
                              'ojTryTime': '', 'optionDtos': '', 'options': '', 'plainTextTitle': '', 'position': '',
                              'sampleAnswerJson': '', 'sampleAnswers': '', 'score': '', 'stdAnswer': '', 'testId': '',
                              'title': '',
                              'titleAttachment': '', 'titleAttachmentDtos': '', 'type': ''}
        num = len(list_0)
        Array = []
        for n_0 in range(num):
            if len(content_id_list) == 2:
                Array.append("reference:c0-e{}".format(6 + 54 * n_0))
            else:
                Array.append("reference:c0-e{}".format(6 + 42 * n_0))

        submit_data["c0-e5"] = "Array:{}".format(Array)
        # Array
        for n in range(num):
            key = Array[n].split(":")[-1]
            num = int(Array[n].split("e")[-1])
            num0 = num

            id = id_list[n]
            plainTextTitle = plainTextTitle_list[n]
            score = score_list[n]
            title = title_list[n]
            type = type_list[n]
            position = position_list[n]

            for key0 in Object_Object_dict.keys():
                if num0 <= num + 16:
                    Object_Object_dict[key0] = "reference:c0-e{}".format(num0 + 1)
                else:
                    Object_Object_dict[key0] = "reference:c0-e{}".format(num0 + 25)
                num0 += 1
            submit_data[key] = "Object_Object:{}".format(Object_Object_dict)
            # Object_Object
            for key_0, value in Object_Object_dict.items():
                key1 = value.split(":")[-1]
                if key_0 == "id":
                    submit_data[key1] = "number:{}".format(id)
                elif key_0 == "plainTextTitle":
                    submit_data[key1] = "string:{}".format(parse.quote(plainTextTitle))
                elif key_0 == "score":
                    submit_data[key1] = "number:{}".format(score.split(".")[0])
                elif key_0 == "title":
                    submit_data[key1] = "string:{}".format(parse.quote(title))
                elif key_0 == "position":
                    submit_data[key1] = "number:{}".format(position)
                elif key_0 == "type":
                    submit_data[key1] = "number:{}".format(type)
                elif key_0 == "optionDtos":
                    option_list0 = content_id_list[n]
                    Array0 = []
                    num1 = int(key1.split('e')[-1]) + 1
                    for i0 in range(len(option_list0)):
                        Array0.append("reference:c0-e{}".format(num1 + 6 * i0))

                    for i1 in range(len(Array0)):
                        content, id = option_list0[i1]
                        key2 = Array0[i1].split(":")[-1]
                        num2 = int(key2.split("e")[-1])
                        num3 = num2 + 1
                        for key3 in option_dict.keys():
                            option_dict[key3] = "reference:c0-e{}".format(num3)
                            num3 += 1

                        for key3, value3 in option_dict.items():
                            key4 = value3.split(":")[-1]
                            if key3 == "id":
                                submit_data[key4] = "number:{}".format(id)
                            elif key3 == "content":
                                submit_data[key4] = "string:{}".format(parse.quote(content))
                            else:
                                submit_data[key4] = "null:null"
                        submit_data[key2] = "Object_Object:{}".format(option_dict)
                    submit_data[key1] = "Array:{}".format(Array0)

                else:
                    submit_data[key1] = "null:null"
                del submit_data[key]
                submit_data[key] = "Object_Object:{}".format(Object_Object_dict)

        del submit_data["c0-e5"]
        submit_data["c0-e5"] = "Array:{}".format(Array)
        end_num = int(list(submit_data.keys())[-4].split("e")[-1]) + 2
        dict0 = {'aid': 'reference:c0-e1', 'answers': 'reference:c0-e2', 'autoSubmit': 'reference:c0-e3',
                 'evaluateVo': 'reference:c0-e4', 'objectiveQList': 'reference:c0-e5', 'showAnalysis': '',
                 'subjectiveQList': '', 'submitStatus': '', 'tid': '', 'tname': '', 'type': ''}
        list0 = ['showAnalysis', 'subjectiveQList', 'submitStatus', 'tid', 'tname', 'type']
        for i in list0:
            key = "c0-e{}".format(end_num)
            if i == 'showAnalysis':
                submit_data[key] = "boolean:true"
            elif i == "subjectiveQList":
                submit_data[key] = "Array:[]"
            elif i == "submitStatus":
                submit_data[key] = "number:1"
            elif i == "tid":
                submit_data[key] = "number:{}".format(tid)
            elif i == "tname":
                submit_data[key] = parse.quote(tname)
            else:
                submit_data[key] = "number:{}".format(end_type)
            dict0[i] = "reference:{}".format(key)
            end_num += 1
        supplement_dict = {
            'c0-param0': "Object_Object:{}".format(dict0),
            'c0-param1': "boolean:false",
            'batchId': self.get_timestamp(),
        }
        submit_data.update(supplement_dict)
        submit_data0 = {}
        for k, j in submit_data.items():
            submit_data0[k.replace(" ", "")] = j.replace("'", "").replace("/", "%2F").replace('https', 'http').replace(
                ' ', '')
        time.sleep(1)
        url = "https://www.icourse163.org/dwr/call/plaincall/MocQuizBean.submitAnswers.dwr"
        self.session.post(url=url, data=submit_data0)
        return aid, tid

    def get_paper(self, aid, tid):
        Answer_list = []
        data = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'MocQuizBean',
            'c0-methodName': 'getQuizPaperDto',
            'c0-id': '0',
            'c0-param0': 'string:{}'.format(tid),
            'c0-param1': 'number:{}'.format(aid),
            'c0-param2': 'boolean:true',
            'batchId': self.get_timestamp(),
        }
        url = "https://www.icourse163.org/dwr/call/plaincall/MocQuizBean.getQuizPaperDto.dwr"
        response = self.session.post(url=url, data=data)
        text = response.text.encode('utf-8').decode('unicode-escape')

        '''题目处理'''
        # 题目信息截取
        title_list = re.findall('title="<p(.*?<)/p>";s\d+\.titleAttachment', text, re.S)
        end_title_list = []
        for title in title_list:
            title = re.sub('<span style=".*?;"  >|</span>|&nbsp;|(&gt;1)', '', title)
            list_0 = re.findall('"http.*?"|>.*?<', title)
            title_end_list = []
            for i in list_0:
                if 'http' in i:
                    list0 = re.findall('"http.*?"|>.*?<', i)
                    for n in range(len(list0)):
                        list0[n] = re.sub('<br >|"|<p>|</p>|</em>|>|<|><"', '', list0[n])
                        if list0[n] != '':
                            title_end_list.append(list0[n])
                else:
                    i = re.sub('<br >|"|<p>|</p>|</em>|>|<', '', i)
                    if i != '':
                        title_end_list.append(i)
            tpl = tuple(title_end_list)
            end_title_list.append(tpl)

        '''选项处理'''
        # 选项截取下来
        option_text_list = re.findall('type.*?allowUpload|type.*?dwr\.engine', text, re.S)
        option_max_list = []
        for option_text in option_text_list:
            option_text = re.sub('<span style=".*?"  >|</span>|&nbsp;|<em style=".*?"  >', '', option_text)
            '''每个题目里的选项截取'''
            # 答案截取
            answer_list = re.findall("s\d+\.answer=(.*?);", option_text)
            str0 = string.ascii_uppercase
            answer0 = ''
            for n0 in range(len(answer_list)):
                if answer_list[n0] == 'true':
                    answer0 += str0[n0]
            Answer_list.append(answer0)
            list1 = re.findall('content="(.*?)";s\d+\.id', option_text)
            option_list = []
            for i in list1:
                i = re.sub('<p style=".*?;"  >', '', i)
                if 'src' in i:
                    list0 = re.findall('"http.*?"|>.*?<', i)
                    for n in range(len(list0)):
                        list0[n] = re.sub('"|<p>|</p>|</em>|>|<', '', list0[n])
                    option_list.append(list0)
                else:
                    i = re.sub('<br >|<p>|</p>|</em>|>|<', '', i)
                    option_list.append(i)
            option_max_list.append(option_list)
        end_option_list = []
        for option_list in option_max_list:
            num = len(option_list)
            str_list = list(string.ascii_uppercase)[:num]
            dic = dict(zip(str_list, option_list))
            end_option_list.append(dic)
        paper_dic = dict(zip(end_title_list, end_option_list))
        return paper_dic, Answer_list

    def word(self, paper_dic, Answer_list, path, test_name):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
        }
        # 试卷document
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        num = 1
        for title_list, option_dic in tqdm(paper_dic.items(), desc=test_name):
            '''开启一个新段落'''
            paragraph = document.add_paragraph('{},'.format(num))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            run = document.paragraphs[-1].add_run()
            for i in title_list:
                if 'http' in i:
                    response = requests.get(url=i, headers=headers)
                    content = response.content
                    with open('core\paper.png', 'wb') as f:
                        f.write(content)
                    run.add_picture('core\paper.png')
                else:
                    try:
                        if i != '':
                            run.add_text(u'{}'.format(i))
                    except:
                        pass

            '''key,为ABCD等,'''
            for key, option_list in option_dic.items():
                '''开启一个新段落,value为option_list'''
                paragraph = document.add_paragraph('{},'.format(key))
                # 设置行距
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.5  # 1.5倍行距
                run = document.paragraphs[-1].add_run()
                for i in option_list:
                    if 'http' in i:
                        response = requests.get(url=i, headers=headers)
                        content = response.content
                        with open('core\paper.png', 'wb') as f:
                            f.write(content)
                        run.add_picture('core\paper.png')
                    else:
                        try:
                            if i != '':
                                run.add_text(u'{}'.format(i))
                        except:
                            pass

            '''开启一个新段落'''
            paragraph = document.add_paragraph('答案:{}'.format(Answer_list[num - 1]))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            num += 1

        document.save('{}/{}_{}.docx'.format(path, test_name, int(time.time())))

    def spider(self):
        b = os.path.exists('data')
        if b == False:
            os.mkdir('data')
        self.get_csrfKey_userId()
        pagessize = self.get_pagesize()
        courses_dic = self.get_courses(pagessize)
        for key, value in courses_dic.items():
            print(key, value[0], value[1])
        print('-' * 100)
        print('''
           ▄██████▄   ▄██████▄          ▄████████  ▄██████▄     ▄████████       ▄█      ███     
          ███    ███ ███    ███        ███    ███ ███    ███   ███    ███      ███  ▀█████████▄ 
          ███    █▀  ███    ███        ███    █▀  ███    ███   ███    ███      ███▌    ▀███▀▀██ 
         ▄███        ███    ███       ▄███▄▄▄     ███    ███  ▄███▄▄▄▄██▀      ███▌     ███   ▀ 
        ▀▀███ ████▄  ███    ███      ▀▀███▀▀▀     ███    ███ ▀▀███▀▀▀▀▀        ███▌     ███     
          ███    ███ ███    ███        ███        ███    ███ ▀███████████      ███      ███     
          ███    ███ ███    ███        ███        ███    ███   ███    ███      ███      ███     
          ████████▀   ▀██████▀         ███         ▀██████▀    ███    ███      █▀      ▄████▀   
                                                               ███    ███                       
        ''')
        while True:
            print('-' * 100)
            key = input("输入非纯字符将会自动退出\n请输入你要选择课程的序号:")
            if key.isdigit() == True:
                path = 'data/{}_{}'.format(courses_dic[key][0], courses_dic[key][1])
                b = os.path.exists(path)
                if b == False:
                    os.mkdir(path)
                courses_id = courses_dic[key][2]
                print('-' * 100)
                print('正在爬取 {} 课程 {} 测验'.format(courses_dic[key][1], courses_dic[key][0]))
                test_name_id_dic = self.get_test_info(courses_id)
                print('开始下载......')
                for test_name, test_id in test_name_id_dic.items():
                    time.sleep(random.randint(2,5))
                    aid, tid = self.submit_paper(test_id)
                    time.sleep(1)
                    paper_dic, Answer_list = self.get_paper(aid, tid)
                    self.word(paper_dic, Answer_list, path, test_name)
            else:
                sys.exit()


m = mooc_spider()
m.spider()
# pyinstaller -F 1.py
